import tkinter as tk
import pyperclip

from docx import Document
from string import punctuation
from tkinter import filedialog
from tkinter import messagebox


def make_modifiers(r):
    form = lambda x, b: x if b else ''
    check = \
        (
            form('*' , any\
                    ((r.italic, r.style.name == 'Emphasis'))),
            form\
                ('**', \
                    r.bold),
            form\
                ('__', \
                    r.underline)
         )
    return ''.join(check)


def create_text(x, c):
    string = '%s'*3 % \
        (f'{c}', x.strip(), f'{c[::-1]} ')
    return string


# def check_paragraphs(doc):
#     base = []
#     for para in doc.paragraphs:
#         new = check_runs(para.runs)
#         base.append(new)
#     return base


def check_paragraphs(doc):
    base = str()
    for para in doc.paragraphs:
        new = check_runs(para.runs)
        base += new + '\n'
    return base


def check_runs(runs, new=str()):
    for r in runs:
        if r.text.isspace():
            continue
        if r.text in punctuation:
            new = new[:-1] + r.text
            continue
        mod = make_modifiers(r)
        new += create_text(r.text, mod)
    return new


def main():
    path = filedialog.askopenfilename()
    doc = Document(path)
    text = check_paragraphs(doc)

    window = tk.Tk()
    window.title('Output')
    window.geometry('650x250')
    output = tk.Label(
                        window,
                        text=text, 
                        font=(14), 
                        wraplength=600, 
                        justify='left'
                     )
    output.grid\
        (row=0, column=0, padx=5, pady=5)
    quit = tk.Button\
        (
            window, 
            text='Quit', 
            command=window.destroy
        )
    quit.grid\
        (row=1, column=0, padx=5, pady=5)
    copy = tk.Button\
        (
            window, 
            text='Copy to Clipboard', 
            command=pyperclip.copy(text)
        )
    copy.grid\
        (row=2, column=0, padx=5, pady=5)
        
    window.mainloop()

    
    
    # print('-', ' -', '  -', sep='\n')
    # print(check_paragraphs(doc))
    # print('  -', ' -', '-', sep='\n')
    #pyperclip.copy()


if __name__ == '__main__':
    main()